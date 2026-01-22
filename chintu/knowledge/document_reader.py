"""
Document Reader for Chintu AI Assistant.

Provides capability to read and understand various document formats:
- PDF files (using pypdf)
- Word documents (using python-docx)
- Text files
- Markdown files

Extracts text with structure and can summarize using LLM.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

logger = logging.getLogger(__name__)

# PDF support
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    logger.warning("pypdf not installed - PDF reading disabled")

# Word document support
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed - DOCX reading disabled")


class DocumentReader:
    """
    Universal document reader supporting multiple formats.
    
    Supports:
    - PDF (.pdf)
    - Word (.docx)
    - Text (.txt)
    - Markdown (.md)
    """
    
    # Supported extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    
    def __init__(self, llm_client=None):
        """
        Initialize document reader.
        
        Args:
            llm_client: Optional LLM client for summarization
        """
        self.llm = llm_client
        
    def can_read(self, file_path: str) -> bool:
        """Check if the file format is supported."""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf' and not HAS_PYPDF:
            return False
        if ext == '.docx' and not HAS_DOCX:
            return False
            
        return ext in self.SUPPORTED_EXTENSIONS
    
    def read_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Read a document file and extract text.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (text_content, metadata)
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return self._read_pdf(path)
        elif ext == '.docx':
            return self._read_docx(path)
        elif ext in ('.txt', '.md', '.markdown'):
            return self._read_text(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _read_pdf(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Read PDF file using pypdf."""
        if not HAS_PYPDF:
            raise ImportError("pypdf is not installed. Run: pip install pypdf")
            
        reader = PdfReader(str(path))
        
        # Extract metadata
        metadata = {
            "format": "pdf",
            "pages": len(reader.pages),
            "filename": path.name,
        }
        
        # Add PDF metadata if available
        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author
            if reader.metadata.subject:
                metadata["subject"] = reader.metadata.subject
        
        # Extract text from all pages
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
        
        full_text = "\n\n".join(text_parts)
        metadata["char_count"] = len(full_text)
        metadata["word_count"] = len(full_text.split())
        
        logger.info(f"Read PDF: {path.name} ({metadata['pages']} pages, {metadata['word_count']} words)")
        
        return full_text, metadata
    
    def _read_docx(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Read Word document using python-docx."""
        if not HAS_DOCX:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
            
        doc = DocxDocument(str(path))
        
        # Extract metadata
        metadata = {
            "format": "docx",
            "filename": path.name,
            "paragraphs": len(doc.paragraphs),
        }
        
        # Extract core properties if available
        if doc.core_properties:
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.subject:
                metadata["subject"] = doc.core_properties.subject
        
        # Extract text with structure
        text_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check if it's a heading
            if para.style and para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    prefix = '#' * level_num
                    text_parts.append(f"\n{prefix} {text}\n")
                except ValueError:
                    text_parts.append(f"\n## {text}\n")
            else:
                text_parts.append(text)
        
        # Also extract tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]")
        
        full_text = "\n".join(text_parts)
        metadata["char_count"] = len(full_text)
        metadata["word_count"] = len(full_text.split())
        
        logger.info(f"Read DOCX: {path.name} ({metadata['paragraphs']} paragraphs, {metadata['word_count']} words)")
        
        return full_text, metadata
    
    def _extract_table(self, table) -> str:
        """Extract text from a Word table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def _read_text(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Read plain text or markdown file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
            
        ext = path.suffix.lower()
        format_type = "markdown" if ext in ('.md', '.markdown') else "text"
        
        metadata = {
            "format": format_type,
            "filename": path.name,
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": len(text.splitlines()),
        }
        
        logger.info(f"Read {format_type.upper()}: {path.name} ({metadata['word_count']} words)")
        
        return text, metadata
    
    def summarize(self, text: str, max_length: int = 500) -> str:
        """
        Summarize document text using LLM.
        
        Args:
            text: Document text to summarize
            max_length: Maximum summary length
            
        Returns:
            Summary text
        """
        if not self.llm:
            # No LLM - return first portion
            if len(text) <= max_length:
                return text
            return text[:max_length] + "..."
            
        # Use LLM for intelligent summary
        prompt = f"""Summarize the following document in a clear, concise manner. 
Focus on the main points and key information.
Keep the summary under {max_length} characters.

DOCUMENT:
{text[:10000]}  # Limit input to avoid token limits

SUMMARY:"""

        try:
            summary = self.llm.generate(prompt)
            return summary.strip()
        except Exception as e:
            logger.warning(f"LLM summarization failed: {e}")
            # Fallback to simple truncation
            if len(text) <= max_length:
                return text
            return text[:max_length] + "..."
    
    def get_key_points(self, text: str, num_points: int = 5) -> List[str]:
        """
        Extract key points from document using LLM.
        
        Args:
            text: Document text
            num_points: Number of key points to extract
            
        Returns:
            List of key points
        """
        if not self.llm:
            # No LLM - return first sentences
            sentences = text.split('.')[:num_points]
            return [s.strip() + '.' for s in sentences if s.strip()]
            
        prompt = f"""Extract the {num_points} most important key points from this document.
Return each point on a new line, starting with a bullet point (-).

DOCUMENT:
{text[:10000]}

KEY POINTS:"""

        try:
            response = self.llm.generate(prompt)
            # Parse bullet points
            lines = response.strip().split('\n')
            points = []
            for line in lines:
                line = line.strip()
                if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                    points.append(line[1:].strip())
                elif line:
                    points.append(line)
            return points[:num_points]
        except Exception as e:
            logger.warning(f"LLM key point extraction failed: {e}")
            return []


# Global instance
_reader: Optional[DocumentReader] = None


def get_document_reader(llm_client=None) -> DocumentReader:
    """Get or create the global document reader."""
    global _reader
    if _reader is None:
        _reader = DocumentReader(llm_client)
    elif llm_client and not _reader.llm:
        _reader.llm = llm_client
    return _reader


def read_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Convenience function to read a document."""
    reader = get_document_reader()
    return reader.read_file(file_path)
